import io
import logging
from typing import Optional
from docx import Document
from fpdf import FPDF

logger = logging.getLogger("services.export_service")


class PDFReport(FPDF):
    """Classe FPDF personnalisee avec en-tete et pied de page soignes."""

    def header(self):
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(16, 185, 129)  # Couleur Émeraude AssIA
        self.cell(0, 10, "AssIA - Procès-Verbal de Réunion", border=False, new_x="LMARGIN", new_y="NEXT", align="L")
        self.set_draw_color(226, 232, 240)
        self.set_line_width(0.5)
        self.line(10, 20, 200, 20)
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}} - Document généré par AssIA", align="C")


def generate_reunion_docx(reunion) -> io.BytesIO:
    """Génère un document Word (.docx) pour le procès-verbal de réunion."""
    doc = Document()

    # Titre principal
    doc.add_heading(reunion.title, level=0)

    status_str = "VALIDE (Officiel)" if getattr(reunion, "status", None) == "valide" else "BROUILLON (En attente de validation)"
    
    # Métadonnées & Structure attendue
    p_meta = doc.add_paragraph()
    p_meta.add_run("Statut : ").bold = True
    p_meta.add_run(f"{status_str}\n")
    
    date_str = reunion.meeting_date.strftime("%d/%m/%Y %H:%M") if reunion.meeting_date else "Non précisé"
    p_meta.add_run("Date : ").bold = True
    p_meta.add_run(f"{date_str}\n")
    
    p_meta.add_run("Participants : ").bold = True
    p_meta.add_run(f"{reunion.participants or 'Non précisé'}\n")

    p_meta.add_run("Objet : ").bold = True
    p_meta.add_run(f"{getattr(reunion, 'objet', None) or 'Non précisé'}\n")

    # Points abordés
    doc.add_heading("Points abordés", level=1)
    doc.add_paragraph(reunion.summary or "Non précisé")

    # Décisions prises
    doc.add_heading("Décisions prises", level=1)
    if reunion.decisions:
        for line in reunion.decisions.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("Aucune décision enregistrée.")

    # Actions à réaliser
    doc.add_heading("Actions à réaliser", level=1)
    if reunion.actions:
        for line in reunion.actions.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("Aucune action à réaliser.")

    # Prochaine réunion
    doc.add_heading("Prochaine réunion", level=1)
    doc.add_paragraph(getattr(reunion, 'prochaine_reunion', None) or "Non précisé")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def clean_text_for_pdf(text: Optional[str]) -> str:
    if not text:
        return ""
    return text.encode("latin-1", "replace").decode("latin-1")


def generate_reunion_pdf(reunion) -> io.BytesIO:
    """Génère un document PDF pour le procès-verbal de réunion."""
    pdf = PDFReport()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Titre
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(15, 23, 42)
    safe_title = clean_text_for_pdf(reunion.title)
    pdf.multi_cell(w=pdf.epw, h=8, txt=safe_title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Métadonnées
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(16, 185, 129)
    pdf.cell(w=pdf.epw, h=6, txt="Informations Générales", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(51, 65, 85)

    status_str = "VALIDÉ (Procès-verbal Officiel)" if getattr(reunion, "status", None) == "valide" else "BROUILLON (En attente de validation)"
    pdf.cell(w=pdf.epw, h=5, txt=f"Statut : {status_str}", new_x="LMARGIN", new_y="NEXT")

    date_str = reunion.meeting_date.strftime("%d/%m/%Y %H:%M") if reunion.meeting_date else "Non précisé"
    pdf.cell(w=pdf.epw, h=5, txt=f"Date : {date_str}", new_x="LMARGIN", new_y="NEXT")

    participants = clean_text_for_pdf(reunion.participants or "Non précisé")
    pdf.multi_cell(w=pdf.epw, h=5, txt=f"Participants : {participants}", new_x="LMARGIN", new_y="NEXT")

    objet = clean_text_for_pdf(getattr(reunion, 'objet', None) or "Non précisé")
    pdf.multi_cell(w=pdf.epw, h=5, txt=f"Objet : {objet}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    def add_section(header: str, content: Optional[str]):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(16, 185, 129)
        pdf.cell(w=pdf.epw, h=7, txt=header, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(51, 65, 85)

        if content:
            safe_content = clean_text_for_pdf(content)
            for line in safe_content.split("\n"):
                line = line.strip()
                if line:
                    if line.startswith("-") or line.startswith("*"):
                        pdf.multi_cell(w=pdf.epw, h=5, txt=f"  {line}", new_x="LMARGIN", new_y="NEXT")
                    else:
                        pdf.multi_cell(w=pdf.epw, h=5, txt=line, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.cell(w=pdf.epw, h=5, txt="Non précisé", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    add_section("Points abordés", reunion.summary)
    add_section("Décisions prises", reunion.decisions)
    add_section("Actions à réaliser", reunion.actions)
    add_section("Prochaine réunion", getattr(reunion, 'prochaine_reunion', None))

    stream = io.BytesIO(pdf.output())
    stream.seek(0)
    return stream
